from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


SRC = Path("ReinforcementSchedules_BlindedManuscript.docx")
OUT = Path("ReinforcementSchedules_BlindedManuscript_methods_revision.docx")


def set_text(paragraph, text, style=None, bold=False):
    paragraph.clear()
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def insert_paragraph_after(paragraph, text="", style=None, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def insert_table_after(paragraph, rows):
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(rows[0]), width=Inches(6.5))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for i, text in enumerate(rows[0]):
        hdr[i].text = text
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    paragraph._p.addnext(table._tbl)
    return table


def find_para(doc, exact=None, starts=None):
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if exact is not None and text == exact:
            return p
        if starts is not None and text.startswith(starts):
            return p
    raise ValueError(f"Paragraph not found: {exact or starts}")


def move_block_before(start_para, end_before_para, target_para):
    elements = []
    el = start_para._p
    stop = end_before_para._p
    while el is not None and el is not stop:
        nxt = el.getnext()
        elements.append(el)
        el = nxt
    for el in elements:
        target_para._p.addprevious(el)


def replace_refs(paragraph, mapping):
    original = paragraph.text
    if not original:
        return
    text = original
    for old, new in mapping.items():
        text = text.replace(old, new)
    if text != original:
        set_text(paragraph, text, paragraph.style)


def add_caption_warning(doc):
    # Keep this as a visible author-facing note only if needed in later passes.
    pass


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    # Give real heading styles to manuscript section labels.
    heading_labels = {
        "The Reinforcement Schedules Application": 1,
        "Prerequisites": 1,
        "The Set-Up Form": 1,
        "The Test Equipment Form": 1,
        "The Component Form": 1,
        "Main Form": 1,
        "The Data Files": 1,
        "Demonstration Experiment": 1,
        "Method": 2,
        "Subjects": 2,
        "Apparatus": 2,
        "Procedure": 2,
        "Data analysis and measures": 2,
        "Results and Discussion": 1,
        "Conclusions": 1,
        "References": 1,
    }
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text in heading_labels:
            p.style = f"Heading {heading_labels[text]}"

    # Move Component Form before Set-Up Form so the implementation workflow is logical.
    component_start = find_para(doc, exact="The Component Form")
    main_start = find_para(doc, exact="Main Form")
    setup_start = find_para(doc, exact="The Set-Up Form")
    move_block_before(component_start, main_start, setup_start)

    # Main argument and tone.
    set_text(
        doc.paragraphs[0],
        "Reinforcement Schedules: An Open-Source Visual Interface for Component-Based Operant Procedures",
        bold=True,
    )
    set_text(
        doc.paragraphs[1],
        "Precise control and recording of experimental events are central requirements in operant research, but many laboratories and teaching programs need tools that are transparent, modifiable, and less costly than commercial control systems. The present paper describes Reinforcement Schedules, an open-source Windows application for configuring, executing, monitoring, and recording component-based operant procedures through a graphical interface. The program extends the Arduino-Visual Basic architecture described in prior low-cost systems by moving routine schedule construction from source-code editing to a visual workflow. Users can define simple, multiple, mixed, and concurrent arrangements across as many as six operandum inputs, assign fixed- and variable-ratio or interval schedules to each input, arrange component-specific stimuli, schedule response feedback and reinforcement delays, randomize component order when appropriate, and save complete session configurations for later reuse. The application communicates with Arduino-compatible firmware through a simple serial protocol: the firmware must report the state of the configured digital inputs and interpret symbolic commands sent by the Visual Basic program to activate or deactivate outputs. A reference sketch supporting six inputs, four stimulus lights, houselight, tone, pellet delivery, and pump control is included with the software. The demonstration experiment illustrates how the system can execute a multi-component resurgence preparation requiring sequential transitions, stimulus changes, changeover delays, and continuous event-level recording. The manuscript emphasizes implementation requirements, data outputs, and the functional contribution of a graphical control layer for reproducible open-source operant research.",
    )

    intro_updates = {
        "Automation and precise control": "Automation and precise control of experimental events have been defining features of the experimental analysis of behavior since its early electromechanical arrangements. Contemporary operant research depends on the ability to arrange contingencies, change stimulus conditions, and record events with enough temporal resolution to reconstruct what occurred during a session. Commercial systems provide this functionality, but their cost, proprietary structure, and dependence on specialized hardware can limit access, modification, and training.",
        "However, across this history": "Open-source and low-cost technologies have therefore become an important complement to commercial systems. Microcontrollers, 3D printing, laser cutting, and user-modifiable software make it possible to build experimental arrangements that are transparent and adaptable while preserving the core requirements of operant research. This approach is especially valuable in teaching laboratories and laboratories with limited equipment budgets, but its relevance is methodological rather than only regional: a useful system should be implementable, inspectable, and modifiable by independent users.",
        "A clear example": "The Arduino-Visual Basic interface described by Escobar and Perez-Herrera (2015) showed that inexpensive Arduino boards could serve as a reliable USB interface between a Windows computer and an operant chamber. Later work extended this open-source infrastructure to 3D-printed chambers for rats and provided detailed hardware designs and example programs (Escobar et al., 2022). Those systems established the feasibility of the hardware architecture, including Arduino-based input and output control, serial communication, and Visual Basic programs for arranging simple or concurrent schedules.",
        "In practice, these requirements": "A remaining barrier is that procedural flexibility often still depends on editing or duplicating source code. Even when the hardware is inexpensive, users who want to change from one schedule arrangement to another, combine procedures across successive components, or prepare variants for teaching exercises may need to understand how the program represents events, timers, schedules, and outputs. For students and researchers without programming experience, that requirement can shift attention away from experimental design and data interpretation toward debugging.",
        "This context underscores": "The present project addresses that barrier by adding a graphical, component-based control layer to the existing Arduino-Visual Basic tradition. Reinforcement Schedules allows users to define components, assign contingencies to multiple operandum inputs, arrange stimuli and delays, save configurations, test equipment, execute sessions, and inspect data files without modifying the source code for routine procedures.",
        "It is conceived": "The contribution is not a replacement for the previously described hardware or firmware approach. Instead, it is a higher-level interface for configuring a wider class of operant procedures within that architecture. This distinction is important: the Arduino remains an input/output interface, whereas the Visual Basic application implements the experimental logic, session structure, schedule progression, event logging, and graphical monitoring. The demonstration experiment was selected to show that the program can execute a procedure with successive components, stimulus-specific transitions, changeover-delay constraints, and a raw event stream suitable for reconstructing within-session behavior.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in intro_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    app_updates = {
        "Reinforcement Schedules is a Windows Forms Application": "Reinforcement Schedules is a Windows Forms application developed in Visual Basic for use with Arduino-compatible microcontroller interfaces. The program preserves the general division of labor of prior Arduino-Visual Basic systems: the Arduino firmware reads physical inputs and activates outputs, whereas the Visual Basic application arranges contingencies, manages session structure, records data, and provides a graphical interface for users. The current version differs from earlier example programs by allowing routine procedures to be configured through forms rather than by editing code.",
        "Broadly, the program organizes": "Broadly, the program organizes procedures as one or more temporal components. Within each component, users can define duration or reinforcer-based termination criteria, number of iterations, component-correlated stimulation, changeover delays, feedback events, reinforcement delays, and independent schedules for each configured operandum. Components may be presented sequentially or, when selected, in a randomized order that preserves the programmed number of presentations. This architecture supports simple schedules, multiple schedules with successive stimulus-correlated components, mixed arrangements in which component order is not signaled, and concurrent arrangements involving more than one available operandum.",
        "The program is distributed": "The program is distributed as open-source software under the GNU General Public License version 3 (GPL-3.0). The repository includes the Visual Basic source code, the compiled application, and a reference Arduino sketch compatible with the serial communication protocol used by the program. Users may use, study, modify, and redistribute the software, provided that redistributed or modified versions preserve the same license and include the corresponding source code. The software source code is freely available at: https://github.com/anonymous-author3/ReinforcementSchedules.",
    }
    last_app = None
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in app_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                last_app = p
                break

    if last_app is not None:
        rows = [
            ["Feature", "Previously described Arduino-Visual Basic examples", "Reinforcement Schedules"],
            [
                "Primary user workflow",
                "Users selected or modified specific Visual Basic programs for simple or two-response concurrent schedules.",
                "Users configure complete sessions through graphical forms and save/load reusable configuration files.",
            ],
            [
                "Procedural scope",
                "Simple schedules and two-operanda concurrent schedules were emphasized.",
                "Simple, multiple, mixed, and concurrent arrangements can be programmed within a component-based session structure.",
            ],
            [
                "Operandum inputs",
                "Example programs focused on one or two responses.",
                "The current architecture supports up to six configured inputs, each with independent schedule, feedback, delay, and reinforcer settings.",
            ],
            [
                "Session structure",
                "Procedures were tied more closely to a specific program file.",
                "Users can combine up to ten components, repeat components, randomize component order, and mix time- and reinforcer-based termination criteria.",
            ],
            [
                "Data output",
                "Programs recorded real-time data for the implemented schedule.",
                "The program writes a continuous time-stamped event file and a summary file suitable for rapid inspection and post-session reconstruction.",
            ],
        ]
        table = insert_table_after(last_app, rows)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        cap = insert_paragraph_after(
            last_app,
            "Table 1. Summary of the main functional differences between previously described Arduino-Visual Basic examples and the present graphical interface.",
        )
        # Move caption before table.
        table._tbl.addprevious(cap._p)

    prereq_updates = {
        "Reinforcement Schedules runs on": "Reinforcement Schedules runs on a standard Windows computer and communicates with the experimental hardware through a USB serial connection. No specialized computer hardware is required. The software has been tested on standard laboratory and classroom computers with a basic processor, 4 GB of RAM, and an available USB port. A single computer can run more than one instance of the application when separate microcontrollers and chambers are assigned to separate serial ports, although the number of simultaneous sessions depends on available system resources.",
        "The software is configured": "The software is configured by default for an operant chamber with multiple digital inputs, stimulus lights, a houselight, an auditory stimulus, and pellet or liquid reinforcement devices. Unused inputs and outputs can be left unassigned, and the software is not species-specific. Any preparation can be used if its response devices and output devices can be represented as digital input states and output commands understood by the microcontroller interface.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in prereq_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    p16 = find_para(doc, starts="The software is configured by default")
    cursor = insert_paragraph_after(p16, "Firmware compatibility and reference sketch", style="Heading 1")
    for text in [
        "Reinforcement Schedules does not require a unique Arduino sketch, but the firmware must satisfy a small set of communication requirements. First, it must continuously read the physical input lines used for responses or sensors and transmit their states to the computer through the serial port in a stable order. Second, it must listen for single-character or otherwise unambiguous symbolic commands sent by the Visual Basic application. Third, it must translate those commands into output changes, such as turning lights on or off, sounding a tone, inserting or retracting response devices, operating a feeder, or activating a pump. Finally, it must return all relevant outputs to a safe inactive state when instructed or when a session ends.",
        "The repository includes a reference sketch, newProgram2026_6inputs_4lights.ino, that implements this protocol for six digital inputs and a set of commonly used chamber outputs. In that sketch, response inputs are read with internal pull-up resistors and reported over serial communication at 9600 baud. The sketch accepts symbolic commands for houselight control, auditory stimulation, four stimulus lights, six actuator outputs associated with operandum devices, pellet delivery through a stepper-driven feeder, pump activation, and a global all-off command. Users may modify pin assignments or output routines to match their own chamber wiring, as long as the order of reported inputs and the symbolic commands expected by the Visual Basic program remain consistent.",
        "This firmware is included to provide a complete, independently implementable pathway from the graphical application to the chamber hardware. At the same time, the firmware is best understood as a reference implementation of an input/output interface rather than as the main contribution of the present paper. Detailed discussions of Arduino-based operant interfaces, relay control, and 3D-printed chamber hardware are available in the prior descriptions of the platform (Escobar & Perez-Herrera, 2015; Escobar et al., 2022).",
    ]:
        cursor = insert_paragraph_after(cursor, text)

    cursor = insert_paragraph_after(cursor, "Implementation workflow", style="Heading 1")
    for text in [
        "A typical session is prepared in five steps. The user first connects the chamber to the Arduino-compatible interface and selects the corresponding serial port. Second, the user opens the Test Equipment Form to verify that each input is detected and that each output operates as expected. Third, the user defines one or more components in the Component Form, assigning schedules, stimuli, feedback, delay parameters, and termination criteria. Fourth, the user reviews the full session in the Set-Up Form, where subject and session identifiers, global timing parameters, component order, and saved configuration files are managed. Fifth, the user starts the session and monitors ongoing behavior in the Main Form while the program writes raw and summary data files.",
        "Describing the workflow in this order separates the experimental logic from the execution interface. The Component Form defines what can happen within a component; the Set-Up Form assembles those components into a session; the Main Form executes the programmed sequence and records what occurred.",
    ]:
        cursor = insert_paragraph_after(cursor, text)

    component_updates = {
        "The Component Form is the core": "The Component Form is the core configuration interface of Reinforcement Schedules because it defines the experimental logic of each component before those components are assembled into a session. Figure 1 shows the standard view of the Component Form. By configuring components independently, users can construct procedures involving simple, multiple, mixed, or concurrent arrangements without writing or modifying source code.",
        "Section A of the Component Form": "Section A of the Component Form (see Figure 1) defines component-level parameters: component name, duration, number of iterations, maximum number of reinforcers, component-correlated stimulation, intermittency of lights or tone, houselight state, and changeover delay (COD). Components may terminate after a fixed duration or after a programmed number of reinforcers. Because this termination rule is stored at the component level, the same session can combine acquisition, transition, and test components with different stopping criteria.",
        "Within each component": "Within each component, schedules can be assigned independently to each configured input (see Figure 1, Section B). The program includes extinction, fixed-ratio, variable-ratio, fixed-interval, and variable-interval schedules. If no schedule is specified for an input, extinction is programmed by default. Variable schedules are generated using a modern implementation of Hantula's (1991) BASIC adaptation of the Fleshler and Hoffman (1962) progression.",
        "The Component Form also allows": "The Component Form also allows users to configure reinforcer type, reinforcer magnitude, and probabilistic mixed-reinforcer arrangements. Reinforcement may involve pellet delivery, liquid delivery, or mixtures in which each reinforcer type occurs with a user-specified probability. Magnitude and delivery duration can be specified separately for each input, allowing asymmetric reinforcement arrangements within the same component.",
        "Contingencies can be further": "Contingencies can be further customized through response feedback and delay parameters. Users may program brief feedback following responses, signaled or unsignaled reinforcement delays, and lever or operandum retraction during delays. These options are important for recurrence and relapse preparations, including resurgence, renewal, reinstatement, and other procedures in which successive components, stimulus changes, worsening transitions, extinction tests, or differential reinforcement histories must be arranged within and across sessions.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in component_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    data_updates = {
        "During session execution": "During session execution, the program generates two data files: a raw continuous event-level record and a session summary file. Data logging begins at session onset and proceeds until the session ends. The raw file is the primary record for reproducible analysis because each line contains elapsed time and an event code, allowing the session to be reconstructed independently of the on-screen display.",
        "On the raw data file": "In the raw data file, response events, reinforcer deliveries, stimulus changes, delay events, component starts and ends, and other programmed events are identified by event codes. Because component boundaries and stimulus changes are explicitly recorded, users can reconstruct cumulative records, response-rate changes around transitions, obtained reinforcer timing, responses during delays, and other moment-to-moment measures directly from the event stream. This design makes the graphical display a monitoring aid rather than the sole source of evidence.",
        "In addition to the raw": "The summary file provides an aggregated overview of session performance, including response counts, reinforcer deliveries, component-level measures, and total session time. It is intended for rapid inspection, teaching, and routine session checks. Detailed analyses should be based on the raw event file, particularly when the goal is to verify temporal relations among responses, stimuli, component transitions, and reinforcer deliveries.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in data_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    demo_updates = {
        "The purpose of the demonstration": "The purpose of the demonstration experiment was to evaluate whether the interface could execute and record a multi-component operant procedure requiring coordinated stimulus transitions, schedule changes, changeover-delay constraints, and continuous event-level recording. The experiment should therefore be read as a functional validation of the software workflow rather than as a new test of resurgence as a behavioral phenomenon.",
        "The experiment was run": "The procedure integrated features of published resurgence preparations, including worsening transitions and component-specific stimulus arrangements. Such preparations are useful software demonstrations because they require a session to change contingencies across successive components while preserving a record that can be analyzed around the transition into the test component.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in demo_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    conclusion_updates = {
        "The present paper introduced": "The present paper introduced Reinforcement Schedules, an open-source graphical interface for arranging component-based operant procedures with Arduino-compatible hardware. The current version provides a complete implementation pathway that includes the Visual Basic application, a reference Arduino sketch, saved configuration files, equipment testing, session execution, and raw event-level data output. Its central contribution is a visual control layer that allows users to build procedures that previously would have required direct code editing.",
        "It should be noted": "The application is not intended to replace commercial experimental control systems or to duplicate the hardware-focused contribution of earlier Arduino-based platforms. Rather, it provides a transparent and modifiable alternative for users who need to configure a wider range of procedures while preserving access to the underlying source code. The system is especially useful for teaching and for laboratories that require adaptable procedures but cannot depend on proprietary software.",
        "Despite its demonstrated": "The current implementation has limitations. It remains optimized for discrete-response operant preparations and for procedures that can be represented as components with scheduled inputs, outputs, delays, and termination criteria. More complex closed-loop adaptive procedures, progressive algorithms, or arrangements requiring continuous analog measurement would require additional development. The present version also depends on accurate hardware wiring and firmware-command compatibility; therefore, users should verify each chamber with the Test Equipment Form before collecting data.",
        "In summary": "In summary, Reinforcement Schedules extends prior low-cost Arduino-Visual Basic systems by making component-based schedule construction available through a graphical interface. By supporting simple, multiple, mixed, and concurrent arrangements across as many as six inputs, saving reusable configurations, and producing raw event files suitable for independent reconstruction, the program provides a practical open-source platform for operant research and instruction.",
    }
    for p in doc.paragraphs:
        compact = " ".join(p.text.split())
        for start, new_text in conclusion_updates.items():
            if compact.startswith(start):
                set_text(p, new_text)
                break

    # Renumber figures affected by moving the Component Form section.
    tmp = {
        "Figure 4": "__FIG_COMP__",
        "Figure 1": "__FIG_SETUP__",
        "Figure 2": "__FIG_CONFIG__",
        "Figure 3": "__FIG_TEST__",
    }
    final = {
        "__FIG_COMP__": "Figure 1",
        "__FIG_SETUP__": "Figure 2",
        "__FIG_CONFIG__": "Figure 3",
        "__FIG_TEST__": "Figure 4",
    }
    for p in doc.paragraphs:
        original = p.text
        if not original:
            continue
        text = original
        for old, new in tmp.items():
            text = text.replace(old, new)
        for old, new in final.items():
            text = text.replace(old, new)
        if text != original:
            set_text(p, text, p.style)

    # The newly rewritten Component Form text already belongs to the first figure after reordering.
    in_component_section = False
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text == "The Component Form" and p.style.name.startswith("Heading"):
            in_component_section = True
        elif text == "The Set-Up Form" and p.style.name.startswith("Heading"):
            in_component_section = False
        elif in_component_section and "Figure 2" in p.text:
            set_text(p, p.text.replace("Figure 2", "Figure 1"), p.style)

    for p in doc.paragraphs:
        if "using the Component Form described below" in p.text:
            set_text(p, p.text.replace("using the Component Form described below", "using the Component Form described above"), p.style)
        if "do not include these comments.." in p.text:
            set_text(p, p.text.replace("do not include these comments..", "do not include these comments."), p.style)

    # Make table numbering later in the manuscript unambiguous after adding Table 1.
    results_seen = False
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text == "Results and Discussion":
            results_seen = True
        elif results_seen and text == "Table 1":
            set_text(p, "Table 2")
            break
        elif results_seen and text.startswith("Table 1 summarizes"):
            set_text(p, text.replace("Table 1 summarizes", "Table 2 summarizes", 1))

    # Caption titles share text with section headings; keep captions out of the heading outline.
    previous_nonempty = ""
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if previous_nonempty.startswith("Figure") and p.style.name.startswith("Heading"):
            p.style = "Normal"
        if text:
            previous_nonempty = text

    # Basic page font consistency for edited paragraphs.
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None and not p.style.name.startswith("Heading"):
                run.font.size = Pt(12)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
